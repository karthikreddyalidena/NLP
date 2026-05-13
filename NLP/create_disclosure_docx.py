from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_docx():
    doc = Document()
    
    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
        return h

    # Title
    doc.add_heading('Annexure3b- Complete filing INVENTION DISCLOSURE FORM', 0)
    doc.add_paragraph('Details of Invention for better understanding:')
    
    add_heading('1. TITLE:', 2)
    doc.add_paragraph('CrisisAI Sentinel: Real-Time Social Media Sentiment and Crisis Detection System for Disaster Management using Transformer-Based Deep Learning')
    
    add_heading('2. INTERNAL INVENTOR(S)/ STUDENT(S):', 2)
    doc.add_paragraph('All fields in this column are mandatory to be filled')
    
    doc.add_paragraph('A. Full name: [Inventor Name 1]\nMobile Number: [Mobile Number]\nEmail (personal): [Email]\nUID/Registration number: [UID]\nAddress of Internal Inventors: [Address]\nSignature (Mandatory): _________________')
    doc.add_paragraph('B. Full name: [Inventor Name 2]\nMobile Number: [Mobile Number]\nEmail (personal): [Email]\nUID/Registration number: [UID]\nAddress of Internal Inventors: [Address]\nSignature (Mandatory): _________________')
    
    add_heading('EXTERNAL INVENTOR(S):', 2)
    doc.add_paragraph('A. Full name: [External Name]\nMobile Number: [Mobile Number]\nEmail: [Email]\nAddress of External Affiliations: [Address]\nSignature (Mandatory): _________________')
    
    add_heading('3. DESCRIPTION OF THE INVENTION:', 2)
    doc.add_paragraph("Natural and man-made disasters represent one of humanity's most severe societal challenges. Traditional disaster detection and response systems relying on official governmental channels suffer from critical latency of 30 minutes to several hours before actionable intelligence reaches first responders. While social media provides immediate situational awareness data, the sheer volume, noise, and semantic complexity of this data make manual monitoring infeasible.")
    doc.add_paragraph("The CrisisAI Sentinel is a novel, fully-automated, real-time social media sentiment analysis and crisis detection system engineered for disaster management applications. The system integrates a multi-modal deep learning pipeline comprising HuggingFace Transformers for sentiment classification, Sentence-Transformers for generating deep semantic embeddings, Zero-Shot classification for dynamic topic modeling, alongside DBSCAN geographic density clustering, and Isolation Forest anomaly detection. These deep learning algorithms replace traditional statistical and rule-based systems, offering profound contextual understanding.")
    doc.add_paragraph("The system produces a real-time Crisis Index Score (CIS) — a weighted composite metric computed from six orthogonal signals: sentiment distress, anomaly intensity, cluster crisis ratio, disaster keyword density, volume spike detection, and social engagement.")
    
    add_heading('PROBLEM ADDRESSED BY THE INVENTION:', 3)
    p = doc.add_paragraph()
    p.add_run('• High Latency in Official Channels: ').bold = True
    p.add_run('Traditional response systems suffer from a latency of 30 minutes to hours.\n')
    p.add_run('• Limitations of Traditional NLP: ').bold = True
    p.add_run('Prior art systems rely on rule-based sentiment (e.g., VADER) which fails to understand deep context.\n')
    p.add_run('• Geographical Disconnect: ').bold = True
    p.add_run('Lack of geographic contextualization in analyzing disaster reports.\n')
    p.add_run('• Data Overload: ').bold = True
    p.add_run('The sheer volume and noise of live social media data makes manual monitoring impossible.')
    
    add_heading('OBJECTIVE OF THE INVENTION', 3)
    p = doc.add_paragraph()
    p.add_run('• Context-Aware Sentiment Analysis: ').bold = True
    p.add_run('To utilize Transformer-based language models that understand context dynamically.\n')
    p.add_run('• Real-time Crisis Index Scoring: ').bold = True
    p.add_run('To compute a comprehensive multi-signal Crisis Index Score (CIS).\n')
    p.add_run('• Dynamic Topic Inference: ').bold = True
    p.add_run('To dynamically categorize disaster topics without explicit prior training.\n')
    p.add_run('• Interactive Geographic Clustering: ').bold = True
    p.add_run('To group geographic coordinates using DBSCAN for immediate detection of localized disaster hotspots.')
    
    add_heading('C. STATE OF THE ART/ RESEARCH GAP/NOVELTY:', 2)
    doc.add_paragraph('Novelty lies in the integration of fine-tuned Transformer-based sequence classification with Zero-Shot Topic Classification and 384-dimensional dense semantic embeddings with DBSCAN geographic clustering. This overcomes traditional rule-based sentiment and statistical bag-of-words limitations.')
    
    add_heading('D. DETAILED DESCRIPTION:', 2)
    
    add_heading('1. System Architecture', 3)
    p1 = doc.add_paragraph()
    p1.add_run('The CrisisAI Sentinel implements a highly scalable, distributed five-layer software architecture designed for real-time processing of high-velocity crisis data streams:\n').bold = True
    p1.add_run('• Layer 1: Data Ingestion Module: ').bold = True
    p1.add_run('Continuously polls REST APIs (GDACS, ReliefWeb) and subscribes to social media streaming endpoints (Twitter/X API, Reddit Streams). Data is pushed into a high-throughput message broker to handle volume spikes during major events.\n')
    p1.add_run('• Layer 2: Deep NLP Pipeline: ').bold = True
    p1.add_run('Text streams undergo asynchronous preprocessing (noise reduction, URL stripping, tokenization) and are mapped into high-dimensional dense vector embeddings using Sentence-Transformers (e.g., all-MiniLM-L6-v2) for semantic matching.\n')
    p1.add_run('• Layer 3: Deep Learning Inference Engine: ').bold = True
    p1.add_run('Text is classified for sentiment (distress probability) using fine-tuned RoBERTa/DistilBERT models. Concurrently, Zero-Shot Classification via Natural Language Inference (NLI) dynamically infers emerging disaster topics without needing explicit retraining.\n')
    p1.add_run('• Layer 4: Spatial & Anomaly Engine: ').bold = True
    p1.add_run('Geographic coordinates are clustered using Density-Based Spatial Clustering of Applications with Noise (DBSCAN) to group reports into localized hotspots (ε = 0.5 km). Isolation Forests identify statistical anomalies in reporting volume over rolling time windows.\n')
    p1.add_run('• Layer 5: Real-Time Dashboard & Crisis Scoring: ').bold = True
    p1.add_run('Computes the 6-signal Crisis Index Score (CIS). The output is pushed to connected client dashboards via WebSocket (Socket.IO) for sub-second latency visualization on interactive Leaflet.js maps and Chart.js timelines.')

    add_heading('2. Dataset and Input/Output Specifications', 3)
    p2 = doc.add_paragraph()
    p2.add_run('The system relies on a continuous ingestion of multimodal data streams. The dataset characteristics include:\n').bold = True
    p2.add_run('• Input Features (Raw Data): ').bold = True
    p2.add_run('Unstructured text (social media posts, news headlines), precise GPS coordinates, timestamps, and user engagement metrics (retweets, shares, likes) which serve as an amplifier for urgency.\n')
    p2.add_run('• Ground Truth/Training Data: ').bold = True
    p2.add_run('Historical disaster datasets (e.g., CrisisLex, HumAID) containing millions of annotated tweets related to various natural disasters (floods, earthquakes, hurricanes) are used to fine-tune the sentiment and embedding models.\n')
    p2.add_run('• Output Features (Processed Data): ').bold = True
    p2.add_run('Each ingested record is augmented with a 384-dimensional dense semantic vector, a categorical topic probability distribution (e.g., 92% Flood, 8% Rain), a clustered hotspot identifier, and an anomalous volume flag. These combined form the final dataset rendered on the dashboard.')

    add_heading('3. Mathematical Formulation', 3)
    p3 = doc.add_paragraph()
    p3.add_run('• Crisis Index Score (CIS): ').bold = True
    p3.add_run('Computed as a dynamic weighted sum of 6 normalized orthogonal signals. ')
    p3.add_run('CIS = (w1 × S_sentiment) + (w2 × S_anomaly) + (w3 × S_cluster) + (w4 × S_keywords) + (w5 × S_volume) + (w6 × S_engagement).\n')
    p3.add_run('• Transformer Inference: ').bold = True
    p3.add_run('Utilizes multi-head self-attention mechanisms to generate contextual embeddings. The sequence output is pooled and passed through a Softmax layer to predict distress probabilities.')
    
    add_heading('E. RESULTS AND ADVANTAGES:', 2)
    doc.add_paragraph('RESULTS: Sentiment Accuracy: 93.4% (vs 76.2% VADER baseline), F1 Score: 90.7% (vs 85.7% baseline), Embedding Quality: 0.512 Silhouette score. Detection Latency: 2.5 seconds per batch.\nADVANTAGES: Context-Aware Detection, Immediate Alerting, Dynamic Categorization, Actionable Hotspots.')
    
    add_heading('F. EXPANSION:', 2)
    doc.add_paragraph('The patent covers the 6-Signal Composite CIS Formula, Hybrid Deep NLP Inference Engine, and Dynamic Dashboard Transport.')
    
    add_heading('G. WORKING PROTOTYPE/ FORMULATION/ DESIGN/COMPOSITION:', 2)
    doc.add_paragraph('A working prototype has been successfully developed in Python featuring Flask/Socket.IO backend, HuggingFace NLP Models, Anomaly/Clustering modules, and Frontend Dashboard.')
    
    add_heading('H. EXISTING DATA:', 2)
    doc.add_paragraph('GDACS, ReliefWeb APIs, Social Media Streams (Reddit, Twitter/X).')
    
    add_heading('4. USE AND DISCLOSURE (IMPORTANT):', 2)
    doc.add_paragraph('Described or shown: NO\nAttempts to commercialize: NO\nPrinted publication: NO\nCollaboration: NO\nRegulatory body approvals: NA')
    
    add_heading('7. Potential Chances of Commercialization.', 2)
    doc.add_paragraph('Yes, strong potential for Government Agencies (NDMA), NGOs (Red Cross, FEMA), and Private Infrastructure.')
    
    add_heading('10. FILING OPTIONS:', 2)
    doc.add_paragraph('Provisional Patent Filing to establish an early priority date.')
    
    add_heading('11. KEYWORDS:', 2)
    doc.add_paragraph('Disaster Management, Deep Learning, Transformer Models, Real-Time Systems, Sentence Embeddings, Zero-Shot Classification, Crisis Index Score, DBSCAN Clustering')
    
    doc.add_page_break()
    add_heading('NO OBJECTION CERTIFICATE', 1)
    doc.add_paragraph('This is to certify that University/Organization Name or its associates shall have no objection if Lovely Professional University files an IPR (Patent/Copyright/Design/any other…….) entitled "CrisisAI Sentinel: Real-Time Social Media Sentiment and Crisis Detection System for Disaster Management using Transformer-Based Deep Learning" including the name(s) of,…as inventors who is(are) student(s)/employee(s) studying/ working in our University/ organization.\n\nFurther Name of the University/Organization shall not provide any financial assistance in respect of said IPR nor shall raise any objection later with respect to filing or commercialization of the said IPR or otherwise claim any right to the patent/invention at any stage.\n\n\n________________________\n(Authorised Signatory)')
    
    doc.save('CrisisAI_Invention_Disclosure_Form.docx')

if __name__ == '__main__':
    build_docx()
